# -*- coding: utf-8 -*-
import json
import os
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"E:\HACHKATON\Hackaton IndabaX Congo"
DONNEES_GENEREES = os.path.join(ROOT, "donnees_generees")
OUT_PATH = os.path.join(ROOT, "Rapport_Methodologique_ACPE_IndabaX2026.docx")
OUT_PATH_PDF = os.path.join(ROOT, "Rapport_Methodologique_ACPE_IndabaX2026.pdf")

with open(os.path.join(DONNEES_GENEREES, "evaluation_final.json"), encoding="utf-8") as f:
    EVAL = json.load(f)

FONT = "Times New Roman"
ACCENT = RGBColor(0x1F, 0x5C, 0x3F)
HEADER_SHADE = "1F5C3F"
LIGHT_SHADE = "EAF2EE"

doc = Document()

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), FONT)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_single_spacing(paragraph, size=10.5, space_after=2):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.name = FONT


def add_h1(text):
    p = doc.add_heading("", level=1)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    r.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_h2(text):
    p = doc.add_heading("", level=2)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(12.5)
    r.font.color.rgb = ACCENT
    r.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_p(text, bold=False, italic=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def add_bullets(items, size=12):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it)
        r.font.name = FONT
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return


def add_table(headers, rows, col_widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(font_size)
        set_cell_shading(hdr_cells[i], HEADER_SHADE)
    for row_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val))
            r.font.name = FONT
            r.font.size = Pt(font_size)
            if row_i % 2 == 1:
                set_cell_shading(cells[i], LIGHT_SHADE)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return t


# ============================================================
# EN-TÊTE (compact, pas de page de garde séparée)
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_after = Pt(2)
r = p.add_run("IndabaX Congo 2026 · ACPE")
r.font.name = FONT
r.bold = True
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Système intelligent d'appariement entre demandeurs d'emploi et offres d'emploi")
r.font.name = FONT
r.font.size = Pt(15)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Rapport méthodologique — Équipe DataUsis (Reine, Daïna, Emmanuelle, Angélica)")
r.font.name = FONT
r.italic = True
r.font.size = Pt(10.5)

# ============================================================
# 1. CONTEXTE ET PROBLEMATIQUE
# ============================================================
add_h1("1. Contexte et problématique")
add_p(
    "L'Agence Congolaise pour l'Emploi (ACPE) reçoit chaque année des milliers d'offres d'emploi et "
    "des dizaines de milliers de profils de demandeurs, mais l'appariement entre les deux reste "
    "aujourd'hui manuel, chronophage et peu traçable. L'objectif poursuivi est la conception d'un "
    "moteur de recommandation Top-K produisant, pour chaque candidat, un classement ordonné d'offres "
    "pertinentes accompagné d'un score explicable, condition posée par le jury pour qu'une solution "
    "soit jugée exploitable par les conseillers de l'Agence."
)
add_p(
    "Un audit intégral des quatre jeux de données fournis a été réalisé avant toute décision de "
    "modélisation, afin d'écarter tout risque de corruption silencieuse des métriques d'évaluation."
)

# ============================================================
# 2. AUDIT DES DONNEES
# ============================================================
add_h1("2. Audit de la qualité des données")
add_p(
    "Quatre fichiers ont été analysés : Offres_ACPE.xlsx (2 535 offres), Offres_ACPE_Extensions.xlsx "
    "(143 offres à texte libre), Demandeurs_.xlsx (41 298 candidats) et "
    "Appariement_Demandeurs_Offres.xlsx (vérité terrain, 3 offres pertinentes par candidat). "
    "Sept anomalies structurelles ont été identifiées et corrigées.",
    space_after=4,
)
add_table(
    ["Anomalie", "Constat", "Décision"],
    [
        ["Référence d'offre dupliquée", "JOB250002109 utilisée par 5 lignes distinctes (Poste 1 à 5)",
         "Désambiguïsation par la clé composite Référence + Poste"],
        ["Doublons de Matricule", "13 Matricules partagés par 2 individus distincts (26 lignes)",
         "Clé composite Matricule + suffixe de ligne"],
        ["Colonne vérité terrain vide", "offre_pertinente nulle à 100 %",
         "Ground truth utilisé exclusivement depuis Appariement_Demandeurs_Offres"],
        ["Secteur d'activité manquant", "9,6 % de valeurs manquantes (3 984/41 298)",
         "Catégorie explicite « Non renseigné », aucune suppression de ligne"],
        ["Mobilité quasi inexploitable", "« Non déclaré » pour 90,6 % des candidats",
         "Département déduit du Matricule utilisé en repli géographique"],
        ["Champ sectoriel ambigu", "« Secteur demandé » ne correspond à l'offre que pour 6,2 % des cas, "
         "contre 86 % pour « secteur_metier »", "Filtrage et score fondés sur secteur_metier"],
        ["Extensions et fichier principal", "Les 143 références sont déjà présentes dans le fichier "
         "principal", "Jointure gauche sur la référence, jamais de concaténation"],
    ],
    col_widths=[3.6, 6.4, 6.5],
)
add_p(
    "Une découverte complémentaire a permis de résoudre l'absence de champ de localisation candidat : "
    "le préfixe du Matricule (par exemple PPBZV, PPPNR, PPKOU) encode le département d'origine et "
    "recoupe les valeurs du champ Lieu des offres, pour l'ensemble des douze départements du Congo.",
    space_after=4,
)

# ============================================================
# 3. ARCHITECTURE DU MOTEUR
# ============================================================
add_h1("3. Architecture du moteur d'appariement")
add_p(
    "Une architecture hybride à trois couches a été retenue afin de répondre à l'exigence "
    "d'explicabilité du jury : le score est décomposé en sous-scores interprétables dès la conception, "
    "plutôt que d'appliquer un outil d'explicabilité a posteriori sur un modèle boîte noire."
)
add_bullets([
    "Secteur et localisation : similarité textuelle entre le secteur_metier du candidat et le secteur "
    "de l'offre, combinée à un score de localisation (département déduit du Matricule et mobilité "
    "déclarée), appliquée de façon progressive pour ne jamais produire de short-list vide.",
    "Texte : similarité TF-IDF sur les intitulés et qualifications, disponible pour 100 % du "
    "catalogue, enrichie par un second TF-IDF sur les champs Description/Profil/Compétences pour "
    "les 143 offres qui en disposent.",
    "Structure : compatibilité entre l'objectif du candidat (Emploi, Stage, Formation) et le groupe "
    "de contrat de l'offre.",
])
add_p(
    "Le score final résulte d'une somme pondérée des trois couches : "
    "Score = w₁ × secteur/localisation + w₂ × texte + w₃ × structure. Les poids ont été calibrés par "
    "recherche d'hyperparamètres sur un jeu de validation de 800 candidats, distinct du jeu de test, "
    "conduisant à w₁ = 0,21, w₂ = 0,70 et w₃ = 0,09.",
    space_after=4,
)

# ============================================================
# 4. PROTOCOLE D'EVALUATION ET PERFORMANCE DU MODELE
# ============================================================
add_h1("4. Protocole d'évaluation et performance du modèle")
add_p(
    "Le split entraînement/validation/test a été réalisé par candidat, jamais par ligne, afin d'éviter "
    "toute fuite d'information via les 3 offres de vérité terrain d'un même candidat. Les performances "
    "sont mesurées par Precision@K, Recall@K et NDCG@K, K = 5 et K = 10.",
    space_after=4,
)
m = EVAL["metrics"]
add_table(
    ["Indicateur", "K = 5", "K = 10"],
    [
        ["Precision@K", f"{m['precision@5']*100:.1f} %", f"{m['precision@10']*100:.1f} %"],
        ["Recall@K", f"{m['recall@5']*100:.1f} %", f"{m['recall@10']*100:.1f} %"],
        ["NDCG@K", f"{m['ndcg@5']:.3f}", f"{m['ndcg@10']:.3f}"],
    ],
    col_widths=[5.5, 5.5, 5.5],
)
add_p(
    f"Résultats obtenus sur un jeu de test de {EVAL['test_size']} candidats, jamais utilisé pour la "
    "calibration des poids. Le ground truth ne contenant que 3 offres pertinentes par candidat, "
    "Precision@5 est plafonnée à 60 % et Precision@10 à 30 % dans l'absolu, indépendamment de la "
    "qualité du modèle ; les valeurs obtenues représentent donc environ la moitié du plafond "
    "théorique à K = 5 et 60 % du plafond à K = 10.",
    space_after=4,
)

# ============================================================
# 5. FONCTIONNALITES COMPLEMENTAIRES
# ============================================================
add_h1("5. Fonctionnalités complémentaires")
add_bullets([
    "Tableau de bord décisionnel : indicateurs macro-économiques, cartographie de tension "
    "emploi/candidats par département, distribution des scores de compatibilité, export PDF.",
    "Recherche en langage naturel : requête libre comparée par similarité cosinus aux offres et aux "
    "candidats, sans correspondance exacte de mots-clés.",
    "Analyse des écarts de compétences (Skill Gap) : compétences citées dans l'offre et absentes du "
    "profil du candidat, disponible pour les 143 offres à texte riche.",
    "Simulateur de compatibilité : estimation en temps réel de l'effet d'un changement de mobilité ou "
    "de l'ajout d'une compétence sur le score, sans modification du profil réel.",
    "Assistant contextuel local : interpréteur des recommandations et indicateurs par détection "
    "d'intention, sans dépendance à une API externe.",
    "Authentification par rôle (Demandeur, Recruteur, Conseiller ACPE), base SQLite, mots de passe "
    "hachés (SHA-256 salé).",
], size=11)
add_p(
    "Stack technique : Python, pandas et scikit-learn pour la préparation des données et le moteur "
    "TF-IDF, Streamlit pour l'interface, Plotly pour les visualisations, ReportLab pour l'export PDF.",
    size=11, space_after=4,
)

# ============================================================
# 6. ORGANISATION DU CODE
# ============================================================
add_h1("6. Organisation du code")
add_p(
    "Le dépôt est structuré en cinq dossiers, séparant les données, la logique du moteur, "
    "l'interface et les données précalculées, conformément à l'exigence d'un dépôt organisé.",
    space_after=4,
)
add_table(
    ["Dossier", "Contenu"],
    [
        ["data/", "Quatre jeux de données bruts fournis par les organisateurs (fichiers Excel), "
         "non modifiés."],
        ["moteur/", "Logique métier : préparation et nettoyage des données (data_prep.py), moteur "
         "d'appariement à 3 couches (matching.py), calibration des poids et évaluation "
         "(calibrate_weights.py, evaluation.py, final_evaluation.py), extraction de compétences "
         "d'un CV (cv_parser.py, skills.py), recherche en langage naturel (search.py), assistant "
         "contextuel (chatbot.py), authentification (auth.py), export PDF (pdf_report.py)."],
        ["application/", "Interface Streamlit : point d'entrée unique (main.py), styles et "
         "composants communs (common.py), une page par profil dans pages/ (login, demandeur, "
         "recruteur, conseiller), images de fond dans images/."],
        ["donnees_generees/", "Données précalculées : recommandations pour l'ensemble des "
         "candidats, statistiques du tableau de bord, résultats d'évaluation, base des comptes "
         "utilisateurs (SQLite)."],
        ["api/", "API REST optionnelle exposant le moteur pour un accès programmatique "
         "(non nécessaire à l'utilisation de l'application)."],
    ],
    col_widths=[3.5, 13.0],
)
add_p(
    "L'application se lance avec la commande streamlit run application/main.py.",
    size=11, space_after=4,
)

# ============================================================
# 7. ACCES DE DEMONSTRATION POUR LE JURY
# ============================================================
add_h1("7. Accès de démonstration pour le jury")
add_p(
    "Un compte a été créé à l'avance pour chacun des trois profils applicatifs : le jury n'a "
    "besoin de créer aucun compte pour tester l'outil. Le lancement de l'application affiche "
    "uniquement la page de connexion ; l'accès aux fonctionnalités de chaque profil s'effectue "
    "directement après authentification avec les identifiants ci-dessous.",
    space_after=4,
)
add_table(
    ["Profil", "Email", "Mot de passe"],
    [
        ["Demandeur d'emploi", "jury.demandeur@acpematcher.cg", "Jury2026Demandeur"],
        ["Offreur d'emploi (Recruteur)", "jury.recruteur@acpematcher.cg", "Jury2026Recruteur"],
        ["Conseiller / Administrateur ACPE", "jury.conseiller@acpematcher.cg", "Jury2026Conseiller"],
    ],
    col_widths=[5.5, 6.5, 4.5],
)

# ============================================================
# 8. POINTS INNOVANTS
# ============================================================
add_h1("8. Points innovants")
add_bullets([
    "Reconstruction d'une donnée absente : aucune colonne de localisation n'existait pour les "
    "candidats ; le département a été déduit du préfixe du Matricule plutôt que d'être ignoré, "
    "ce qui a rendu possible la couche secteur/localisation pour l'ensemble des candidats.",
    "Explicabilité par construction : le score n'est pas une boîte noire expliquée a posteriori, "
    "il est nativement décomposé en trois sous-scores interprétables (secteur/localisation, "
    "texte, structure), directement restitués à l'utilisateur.",
    "Pondération calibrée empiriquement : les poids du modèle hybride résultent d'une recherche "
    "d'hyperparamètres comparant règles seules, texte seul et hybride sur un jeu de validation, "
    "plutôt que d'un choix arbitraire.",
    "Accès élargi au-delà du dataset fourni : un candidat absent de la base ACPE peut créer un "
    "profil en téléversant simplement son CV, dont les compétences sont extraites automatiquement "
    "et injectées dans le même moteur de recommandation.",
    "Simulateur d'impact « et si ? » : le candidat peut estimer, avant d'agir, l'effet d'une "
    "compétence à acquérir ou d'un changement de mobilité sur son score de compatibilité avec une "
    "offre donnée, transformant le moteur en outil d'orientation proactif plutôt que descriptif.",
], size=11)

doc.save(OUT_PATH)
print("Saved to", OUT_PATH)

try:
    from docx2pdf import convert
    convert(OUT_PATH, OUT_PATH_PDF)
    print("Saved to", OUT_PATH_PDF)
except Exception as e:
    print("Conversion PDF non effectuée automatiquement :", e)
